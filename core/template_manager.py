"""
Template Manager for AI Builder
Manages document templates and export functionality
"""

import os
import json
from pathlib import Path
from typing import Dict, Any, List, Optional
from jinja2 import Environment, FileSystemLoader, Template
from core.logger import logger
from core.config_manager import config


class TemplateManager:
    def __init__(self):
        self.template_dir = Path(__file__).parent.parent / "templates"
        self.export_formats = ["md", "html", "pdf", "docx", "json"]
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            trim_blocks=True,
            lstrip_blocks=True
        )
        self._setup_template_directory()
    
    def _setup_template_directory(self):
        """Ensure template directory exists"""
        try:
            os.makedirs(self.template_dir, exist_ok=True)
            logger.logger.info(f"Template directory ready: {self.template_dir}")
        except Exception as e:
            logger.log_error(e, "Setting up template directory")
    
    def get_available_templates(self) -> Dict[str, List[str]]:
        """Get list of available templates by category"""
        templates = {
            "documents": [],
            "reports": [],
            "features": [],
            "analysis": []
        }
        
        try:
            for file_path in self.template_dir.rglob("*.md"):
                template_name = file_path.stem
                category = self._categorize_template(template_name)
                templates[category].append(template_name)
            
            logger.logger.info(f"Found templates: {sum(len(v) for v in templates.values())}")
            return templates
            
        except Exception as e:
            logger.log_error(e, "Getting available templates")
            return templates
    
    def _categorize_template(self, template_name: str) -> str:
        """Categorize template based on name"""
        if any(doc in template_name.lower() for doc in ["brd", "srs", "document"]):
            return "documents"
        elif any(feat in template_name.lower() for feat in ["feature", "spec"]):
            return "features"
        elif any(analysis in template_name.lower() for analysis in ["analysis", "architecture"]):
            return "analysis"
        else:
            return "reports"
    
    def load_template(self, template_name: str) -> Optional[Template]:
        """Load a specific template"""
        try:
            template_file = f"{template_name}.md"
            template = self.jinja_env.get_template(template_file)
            logger.logger.info(f"Template loaded: {template_name}")
            return template
            
        except Exception as e:
            logger.log_error(e, f"Loading template: {template_name}")
            return None
    
    def render_template(self, template_name: str, data: Dict[str, Any]) -> Optional[str]:
        """Render template with provided data"""
        try:
            template = self.load_template(template_name)
            if not template:
                return None
            
            rendered_content = template.render(**data)
            logger.logger.info(f"Template rendered successfully: {template_name}")
            return rendered_content
            
        except Exception as e:
            logger.log_error(e, f"Rendering template: {template_name}")
            return None
    
    def export_document(self, content: str, output_path: str, format_type: str = "md", 
                       template_name: str = None, metadata: Dict[str, Any] = None) -> str:
        """Export document in specified format using template"""
        try:
            if format_type not in self.export_formats:
                raise ValueError(f"Unsupported format: {format_type}. Supported: {self.export_formats}")
            
            # Ensure output directory exists
            output_dir = Path(output_path).parent
            os.makedirs(output_dir, exist_ok=True)
            
            # Apply template if specified
            if template_name:
                template_data = {
                    "content": content,
                    "metadata": metadata or {},
                    "timestamp": config.get_current_timestamp(),
                    "version": config.output.current_version
                }
                content = self.render_template(template_name, template_data) or content
            
            # Export based on format
            exported_path = self._export_by_format(content, output_path, format_type, metadata)
            
            logger.logger.info(f"Document exported: {exported_path}")
            return exported_path
            
        except Exception as e:
            logger.log_error(e, f"Exporting document to: {output_path}")
            raise
    
    def _export_by_format(self, content: str, output_path: str, format_type: str, 
                         metadata: Dict[str, Any] = None) -> str:
        """Export content in specific format"""
        base_path = Path(output_path).with_suffix(f".{format_type}")
        
        if format_type == "md":
            return self._export_markdown(content, str(base_path))
        elif format_type == "html":
            return self._export_html(content, str(base_path), metadata)
        elif format_type == "pdf":
            return self._export_pdf(content, str(base_path), metadata)
        elif format_type == "docx":
            return self._export_docx(content, str(base_path), metadata)
        elif format_type == "json":
            return self._export_json(content, str(base_path), metadata)
        else:
            raise ValueError(f"Format {format_type} not implemented")
    
    def _export_markdown(self, content: str, output_path: str) -> str:
        """Export as Markdown"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return output_path
    
    def _export_html(self, content: str, output_path: str, metadata: Dict[str, Any] = None) -> str:
        """Export as HTML"""
        try:
            import markdown
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content, extensions=['tables', 'toc'])
            
            # Wrap in HTML template
            html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{metadata.get('title', 'AI Builder Document') if metadata else 'AI Builder Document'}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_template)
            return output_path
            
        except ImportError:
            logger.logger.warning("markdown package not available, falling back to plain text HTML")
            html_content = f"<html><body><pre>{content}</pre></body></html>"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return output_path
    
    def _export_pdf(self, content: str, output_path: str, metadata: Dict[str, Any] = None) -> str:
        """Export as PDF"""
        try:
            # Try using markdown2pdf or similar
            html_path = output_path.replace('.pdf', '.html')
            self._export_html(content, html_path, metadata)
            
            # Convert HTML to PDF (requires additional packages)
            logger.logger.warning("PDF export requires additional packages (weasyprint or similar)")
            logger.logger.info(f"HTML version available at: {html_path}")
            return html_path
            
        except Exception as e:
            logger.log_error(e, "PDF export")
            # Fallback to text export
            txt_path = output_path.replace('.pdf', '.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return txt_path
    
    def _export_docx(self, content: str, output_path: str, metadata: Dict[str, Any] = None) -> str:
        """Export as DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title if metadata available
            if metadata and 'title' in metadata:
                title = doc.add_heading(metadata['title'], 0)
            
            # Add content (basic markdown parsing)
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line:
                    doc.add_paragraph(line)
            
            doc.save(output_path)
            return output_path
            
        except ImportError:
            logger.logger.warning("python-docx package not available, falling back to text")
            txt_path = output_path.replace('.docx', '.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return txt_path
    
    def _export_json(self, content: str, output_path: str, metadata: Dict[str, Any] = None) -> str:
        """Export as JSON"""
        json_data = {
            "content": content,
            "metadata": metadata or {},
            "exported_at": config.get_current_timestamp(),
            "version": config.output.current_version
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        return output_path
    
    def create_template(self, template_name: str, template_content: str, category: str = "custom") -> str:
        """Create a new template"""
        try:
            template_path = self.template_dir / f"{template_name}.md"
            
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write(template_content)
            
            logger.logger.info(f"Template created: {template_name}")
            return str(template_path)
            
        except Exception as e:
            logger.log_error(e, f"Creating template: {template_name}")
            raise
    
    def export_with_template(self, data: Dict[str, Any], template_name: str, 
                           output_path: str, format_type: str = "md") -> str:
        """Export data using specified template and format"""
        try:
            # Render content using template
            content = self.render_template(template_name, data)
            if not content:
                raise ValueError(f"Failed to render template: {template_name}")
            
            # Export in specified format
            exported_path = self.export_document(
                content=content,
                output_path=output_path,
                format_type=format_type,
                metadata=data.get('metadata', {})
            )
            
            logger.logger.info(f"Exported with template {template_name}: {exported_path}")
            return exported_path
            
        except Exception as e:
            logger.log_error(e, f"Export with template: {template_name}")
            raise
    
    def get_template_variables(self, template_name: str) -> List[str]:
        """Extract variables from template for validation"""
        try:
            template = self.load_template(template_name)
            if not template:
                return []
            
            # Extract Jinja2 variables
            variables = list(template.environment.parse(template.source).find_all(
                lambda node: hasattr(node, 'name')
            ))
            
            var_names = [var.name for var in variables if hasattr(var, 'name')]
            return list(set(var_names))  # Remove duplicates
            
        except Exception as e:
            logger.log_error(e, f"Getting template variables: {template_name}")
            return []


# Create template manager instance
template_manager = TemplateManager()
